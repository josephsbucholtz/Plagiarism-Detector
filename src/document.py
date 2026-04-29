from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import library_manager
import text_processing


def build_highlighted_doc(
    words: list[str],
    matching_shingles: set[str],
    file1_name: str,
    file2_name: str,
    similarity: float,
    output_name: str,
    shingle_size: int,
) -> str:
    library_manager.ensure_runtime_directories()
    doc = Document()

    header = f"HIGHLIGHT MAP: {file1_name} compared to {file2_name}"
    paragraph = doc.add_paragraph()
    title_run = paragraph.add_run(header)
    title_run.bold = True
    paragraph.add_run(f" ({similarity:.2f}% similar)")
    content_paragraph = doc.add_paragraph()

    index = 0
    while index < len(words):
        candidate_words = words[index:index + shingle_size]
        candidate_shingle = " ".join(candidate_words)

        if len(candidate_words) == shingle_size and candidate_shingle in matching_shingles:
            run = content_paragraph.add_run(candidate_shingle + " ")
            run.font.highlight_color = WD_COLOR_INDEX.RED
            index += shingle_size
            continue

        content_paragraph.add_run(words[index] + " ")
        index += 1

    full_path = library_manager.HIGHLIGHT_DIR / output_name
    doc.save(full_path)
    return str(full_path)


def create_highlight_documents(
    file1: str,
    file2: str,
    similarity_percent: float,
    include_second_document: bool = False,
) -> list[str]:
    profile1 = library_manager.build_document_profile(file1)
    profile2 = library_manager.build_document_profile(file2)
    common_shingles = set(profile1["shingles"]).intersection(profile2["shingles"])
    if not common_shingles:
        return []

    file1_name = Path(file1).stem
    file2_name = Path(file2).stem
    shingle_size = profile1.get("shingle_size", library_manager.DEFAULT_SHINGLE_SIZE)
    output_paths = [
        build_highlighted_doc(
            profile1["normalized_text"].split(),
            common_shingles,
            file1_name,
            file2_name,
            similarity_percent,
            f"highlight-{file1_name}-{file2_name}.docx",
            shingle_size,
        )
    ]

    if include_second_document:
        output_paths.append(
            build_highlighted_doc(
                profile2["normalized_text"].split(),
                common_shingles,
                file2_name,
                file1_name,
                similarity_percent,
                f"highlight-{file2_name}-{file1_name}.docx",
                shingle_size,
            )
        )

    return output_paths


def clean_document_data(input_file: str):
    return text_processing.clean_document_data(input_file, use_lemma=False)


def lemma_clean_document_data(input_file: str):
    return text_processing.clean_document_data(input_file, use_lemma=True)


def get_wordmap_documents(file1: str, file2: str, similarity: float):
    return create_highlight_documents(file1, file2, similarity, include_second_document=True)


def get_wordmap_doc(file1: str, file2: str, similarity: float):
    output_paths = create_highlight_documents(file1, file2, similarity, include_second_document=False)
    return output_paths[0] if output_paths else None
